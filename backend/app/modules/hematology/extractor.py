"""
Extrae valores CBC y metadatos del paciente desde PDFs, CSV, Excel e imagenes.

Formato primario PDF: IDEXX ProCyte One (pares de paginas: resultados + scatter/comentarios).
Alternativa PDF: escaneo generico de nombres de campo CBC en cualquier pagina.
Imagenes (JPG, PNG, TIFF, WEBP): OCR con pytesseract + preprocesamiento Pillow.

API publica
-----------
extract_from_pdf(pdf_bytes: bytes) -> ExtractionResult
extract_from_file(contents: bytes, content_type: str) -> ExtractionResult
    Lanza ExtractionError si no se encuentran valores CBC.
"""

from __future__ import annotations

import io
import re
from typing import Optional

import pdfplumber
from app.modules.gemini_extraction.normalizer import (
    canonical_model_key,
    normalize_extracted_payload,
    normalize_label,
)
from app.modules.gemini_extraction.utils.file_text_extraction import decode_text
from .extraction_types import (
    REQUIRED_FIELDS,
    ExtractionError,
    ExtractionResult,
    coerce_lab_number as _coerce_lab_number,
)

# ---------------------------------------------------------------------------
# Definicion de campos (replica la lista CBC_FIELDS de NB01)
# ---------------------------------------------------------------------------

# Lista ordenada de nombres de campo CBC tal como aparecen en los reportes IDEXX.
# Los campos de porcentaje tienen el prefijo "% " en el texto fuente.
CBC_FIELDS = [
    "RBC",
    "Hematocrit",
    "Hemoglobin",
    "MCV",
    "MCHC",
    "MCH",
    "RDW",
    "% Reticulocytes",
    "Reticulocytes",
    "WBC",
    "% Neutrophils",
    "% Lymphocytes",
    "% Monocytes",
    "% Eosinophils",
    "% Basophils",
    "Neutrophils",
    "Lymphocytes",
    "Monocytes",
    "Eosinophils",
    "Basophils",
    "Platelets",
    "PDW",
    "MPV",
    "Plateletcrit",
]

# Mapeo de nombres crudos de extraccion a nombres canonicos del modelo.
# Claves: tal como se extraen (tras normalizacion pct_). Valores: nombres de columna del modelo.
RENAME_TO_CANONICAL = {
    "Hemoglobin": "HGB",
    "Hematocrit": "HCT",
    "Plateletcrit": "PCT",
    "pct_Reticulocytes": "Reticulocytes_pct",
    "pct_Neutrophils": "Neutrophils_pct",
    "pct_Lymphocytes": "Lymphocytes_pct",
    "pct_Monocytes": "Monocytes_pct",
    "pct_Eosinophils": "Eosinophils_pct",
    "pct_Basophils": "Basophils_pct",
}

# Patrones regex para el encabezado del paciente en reportes IDEXX.
HEADER_PATTERNS = {
    "patient_name": r"^(.+?)\nPET OWNER:",
    "pet_owner": r"PET OWNER:\s*(.+?)\s*(?:VET|LAB)",
    "clinic": r"VET\s+(.+?)\s*LAB ID",
    "species": r"SPECIES:\s*(\w+)",
    "breed": r"BREED:\s*(.+?)\s*(?:santo|ORDER|$)",
    "gender": r"GENDER:\s*(\w+)",
    "age": r"AGE:\s*(\d[^\n]+?)\s*(?:ACCOUNT|\n|$)",
    "date_receipt": r"DATE\s*OF RECEIPT:\s*(\d+/\d+/\d+)",
    "date_result": r"DATE OF RESULT:\s*(\d+/\d+/\d+)",
    "attending_vet": r"ATTENDING VET:\s*(.+?)$",
}

# Patrones genericos de campo CBC para PDFs no-IDEXX (abreviaciones comunes).
_GENERIC_VALUE_PATTERN = (
    r"([*<>\u2264\u2265]?\s*[-+]?(?:\d+(?:[.,]\d*)?|[.,]\d+)(?:[eE][-+]?\d+)?\s*\*?)"
)
GENERIC_PATTERNS: dict[str, str] = {
    "WBC": rf"WBC\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "RBC": rf"RBC\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "HGB": rf"(?:HGB|HB|Hemoglobin)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "HCT": rf"(?:HCT|HT|Hematocrit|PCV)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "MCV": rf"MCV\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "MCH": rf"MCH\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "MCHC": rf"MCHC\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "RDW": rf"RDW\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "Platelets": rf"(?:Platelets?|PLT)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "MPV": rf"MPV\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "PDW": rf"PDW\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "Neutrophils": rf"(?:Neutrophils?|NEU|NEUT)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "Lymphocytes": rf"(?:Lymphocytes?|LYM|LYMPH)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "Monocytes": rf"(?:Monocytes?|MONO)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "Eosinophils": rf"(?:Eosinophils?|EOS)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
    "Basophils": rf"(?:Basophils?|BASO)\s*[:\-]?\s*{_GENERIC_VALUE_PATTERN}",
}

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _extract_header(text: str) -> dict[str, Optional[str]]:
    """Extrae metadatos del paciente y clinica del encabezado de una pagina de resultados IDEXX."""
    header: dict[str, Optional[str]] = {}
    for field_name, pattern in HEADER_PATTERNS.items():
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        header[field_name] = match.group(1).strip() if match else None
    return header


def _extract_cbc_idexx(text: str) -> dict[str, float]:
    """
    Extrae valores CBC de una pagina de resultados IDEXX.

    Cada campo CBC ocupa su propia linea iniciando con el nombre del campo.
    El valor numerico sigue inmediatamente (puede tener prefijo * para valores fuera de rango).
    """
    values: dict[str, float] = {}

    for field_name in CBC_FIELDS:
        for line in text.split("\n"):
            line = line.strip()
            if not line.startswith(field_name):
                continue
            remainder = line[len(field_name) :].strip()
            parsed_value = _coerce_lab_number(remainder)
            if parsed_value is not None:
                raw_name = field_name.replace("% ", "pct_").replace(" ", "_")
                values[raw_name] = parsed_value
            break

    return values


def _extract_comments_idexx(text: str) -> Optional[str]:
    """
    Extrae comentarios del analizador IDEXX de una pagina de scatter plots.

    Los comentarios aparecen despues del segundo marcador 'Download' de la pagina.
    Replica la logica de extraccion de NB01.
    """
    if not text:
        return None

    lines = text.split("\n")
    comment_lines: list[str] = []
    download_count = 0

    for line in lines:
        stripped = line.strip()
        if stripped == "Download":
            download_count += 1
            continue
        if "Generated by" in stripped or "IDEXX Laboratories" in stripped:
            break
        if download_count >= 2 and stripped:
            comment_lines.append(stripped)

    comment = " ".join(comment_lines).strip().rstrip("®").strip()
    return comment if comment else None


def _parse_age_years(age_str: Optional[str]) -> Optional[float]:
    """
    Convierte una cadena de edad IDEXX a anios decimales.

    Ejemplos:
        "3 Years 2 Months" -> 3.17
        "18 Months"        -> 1.5
        "5 Years"          -> 5.0
        "6"                -> 6.0
    """
    if not age_str:
        return None

    years = 0.0
    months = 0.0

    y_match = re.search(r"(\d+)\s*(?:year|ano|ano)", age_str, re.IGNORECASE)
    m_match = re.search(r"(\d+)\s*(?:month|mes)", age_str, re.IGNORECASE)

    if y_match:
        years = float(y_match.group(1))
    if m_match:
        months = float(m_match.group(1))

    if not y_match and not m_match:
        # Intentar numero simple (se asume que son anios)
        bare = re.match(r"^(\d+(?:\.\d+)?)", age_str.strip())
        if bare:
            return float(bare.group(1))
        return None

    return round(years + months / 12.0, 2)


def _extract_location(text: str, clinic: Optional[str]) -> Optional[str]:
    """
    Intenta extraer una ubicacion del texto del PDF.

    Estrategia:
    1. Usa el nombre de la clinica del encabezado IDEXX si esta disponible.
    2. Busca nombres de ciudades conocidas de Republica Dominicana como alternativa.
    """
    if clinic:
        return clinic

    # Ciudades RD comunes usadas como escaneo de ubicacion alternativo
    cities = [
        "Santo Domingo",
        "Santiago",
        "La Romana",
        "San Pedro de Macoris",
        "Puerto Plata",
        "Higuey",
        "San Cristobal",
        "La Vega",
        "Moca",
        "Bonao",
        "Bani",
        "Azua",
        "Barahona",
        "Monte Cristi",
        "Samana",
        "Cotui",
    ]
    for city in cities:
        if re.search(re.escape(city), text, re.IGNORECASE):
            return city

    return None


def _rename_to_canonical(raw: dict[str, float]) -> dict[str, float]:
    """Aplica el mapeo RENAME_TO_CANONICAL a los nombres de campo extraidos."""
    canonical: dict[str, float] = {}
    for key, val in raw.items():
        canonical[RENAME_TO_CANONICAL.get(key, key)] = val
    return canonical


def _extract_cbc_generic(full_text: str) -> dict[str, float]:
    """
    Extractor de respaldo para PDFs no-IDEXX.

    Escanea el texto completo buscando abreviaciones CBC comunes y extrae
    el primer valor numerico que sigue a cada patron.
    """
    values: dict[str, float] = {}
    for field_name, pattern in GENERIC_PATTERNS.items():
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            parsed_value = _coerce_lab_number(match.group(1))
            if parsed_value is not None:
                values[field_name] = parsed_value
    return values


_TEXT_VALUE_PATTERN = re.compile(
    r"([*<>\u2264\u2265]?\s*[-+]?(?:\d+(?:[.,]\d*)?|[.,]\d+)(?:[eE][-+]?\d+)?\s*\*?)"
)
_PAREN_CODE_RE = re.compile(r"\(([A-Za-z0-9/#%._-]{2,16})\)")
_LEADING_ITEM_RE = re.compile(r"^\s*\d+(?:-\d+)?[.)]\s*")
_UNUSED_SUBPARAMETER_MARKERS = {
    "nrbc",
    "nst",
    "nsg",
    "nsh",
    "slym",
    "llym",
    "awbc",
    "hdw",
    "etg",
    "sph",
    "aca",
    "agg",
    "heb",
    "aplt",
    "lplt",
    "atipicos",
    "cant",
}
_INDEXED_FORM1_FIELDS = {
    1: "WBC",
    2: "Lymphocytes",
    3: "Monocytes",
    4: "Neutrophils",
    5: "Lymphocytes_pct",
    6: "Monocytes_pct",
    7: "Neutrophils_pct",
    8: "RBC",
    9: "HGB",
    10: "HCT",
    11: "MCV",
    12: "MCH",
    13: "MCHC",
    14: "RDW",
    16: "Platelets",
    17: "MPV",
    18: "PDW",
}
_INDEXED_ROW_RE = re.compile(r"^[^A-Za-z0-9]*(\d{1,2})\s+(.+)$")


def _line_looks_non_result(line: str) -> bool:
    folded = normalize_label(line)
    ignored = ("curve", "threshold", "thresholds", "alarm", "alarms", "interpretive")
    tokens = set(folded.split())
    return bool(
        tokens.intersection(ignored)
        or tokens.intersection(_UNUSED_SUBPARAMETER_MARKERS)
    )


def _first_result_number_after_label(line: str, label_end: int = 0) -> float | None:
    tail = line[label_end:].strip()
    matches = list(_TEXT_VALUE_PATTERN.finditer(tail))
    if not matches:
        return None
    return _coerce_lab_number(matches[0].group(1))


def _normalize_indexed_value(key: str, value: float) -> float:
    if key == "HGB":
        if value > 300:
            return value / 100.0
        if value > 40:
            return value / 10.0
    if key in {"MCV", "MCH", "MCHC"} and value > 150:
        return value / 10.0 if value < 600 else value / 100.0
    return value


def _parse_indexed_form1_line(line: str) -> tuple[str, float] | None:
    if re.match(r"^\s*\d+\s*-\s*\d+", line):
        return None
    match = _INDEXED_ROW_RE.match(line)
    if not match:
        return None
    index = int(match.group(1))
    key = _INDEXED_FORM1_FIELDS.get(index)
    if key is None:
        return None
    value = _first_result_number_after_label(match.group(2))
    if value is None:
        return None
    return key, _normalize_indexed_value(key, value)


def _extract_cbc_from_text_lines(text: str) -> dict[str, float]:
    """
    Conservative parser for OCR/plain-text hemograms.

    It detects explicit aliases and uses the first numeric token after the label,
    which avoids common reference-range columns later in the same line.
    """
    raw_values: dict[str, float] = {}
    for raw_line in text.splitlines():
        indexed = _parse_indexed_form1_line(raw_line.strip())
        if indexed is not None:
            key, value = indexed
            raw_values[key] = value
            continue

        line = _LEADING_ITEM_RE.sub("", raw_line.strip())
        if not line or _line_looks_non_result(line):
            continue

        key = canonical_model_key(line)
        if key is None:
            continue

        # Match either a parenthesized analyzer code or the beginning label.
        value = None
        paren_match = _PAREN_CODE_RE.search(line)
        if paren_match:
            value = _first_result_number_after_label(line, paren_match.end())
        if value is None:
            value = _first_result_number_after_label(line)
        if value is not None:
            folded = normalize_label(line)
            if key in {"HGB", "MCHC"} and "g l" in folded and "g dl" not in folded:
                value = value / 10.0
            raw_values[line] = value

    return normalize_extracted_payload(raw_values).normalized_data


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract_from_pdf(pdf_bytes: bytes) -> ExtractionResult:
    """
    Extrae valores CBC, metadatos del paciente y comentarios de un PDF de hemograma.

    Intenta el formato IDEXX ProCyte One primero (iteracion por pares de paginas).
    Si no se detecta la estructura IDEXX, cae en escaneo generico de texto.

    Parametros
    ----------
    pdf_bytes : bytes
        Contenido crudo del archivo PDF.

    Retorna
    -------
    ExtractionResult
        Poblado con valores CBC (nombres canonicos), metadatos y opcionalmente
        comentarios del analizador IDEXX.

    Lanza
    -----
    ExtractionError
        Si no se encuentran suficientes valores CBC en el PDF.
    """
    try:
        pdf = pdfplumber.open(io.BytesIO(pdf_bytes))
    except Exception as exc:
        raise ExtractionError(f"No se pudo leer el PDF: {exc}") from exc

    pages = pdf.pages
    n_pages = len(pages)

    cbc_raw: dict[str, float] = {}
    header: dict[str, Optional[str]] = {}
    comments: Optional[str] = None
    idexx_found = False

    # Extraccion primaria IDEXX: pares de paginas (resultados + scatter)
    for i in range(0, n_pages, 2):
        try:
            text_results = pages[i].extract_text() or ""
        except Exception:
            continue

        # Las paginas de resultados IDEXX contienen la palabra "Hematology"
        if "Hematology" not in text_results:
            continue

        page_cbc = _extract_cbc_idexx(text_results)
        if not page_cbc:
            continue

        # Tomar solo el primer paciente encontrado en el PDF
        if not idexx_found:
            header = _extract_header(text_results)
            cbc_raw = page_cbc
            idexx_found = True

            # Pagina complementaria de scatter/comentarios
            if i + 1 < n_pages:
                try:
                    text_scatter = pages[i + 1].extract_text() or ""
                    comments = _extract_comments_idexx(text_scatter)
                except Exception:
                    comments = None

            break  # Solo se toma el primer paciente por subida

    pdf.close()

    # Extraccion generica como respaldo
    generic_full_text = ""
    if not idexx_found or not cbc_raw:
        try:
            pdf2 = pdfplumber.open(io.BytesIO(pdf_bytes))
            generic_full_text = "\n".join((p.extract_text() or "") for p in pdf2.pages)
            pdf2.close()
        except Exception as exc:
            raise ExtractionError(f"No se pudo leer el PDF: {exc}") from exc

        cbc_raw = _extract_cbc_generic(generic_full_text)
        header = {}  # sin encabezado estructurado en modo generico

    # Renombrar nombres crudos a nombres canonicos del modelo
    cbc = _rename_to_canonical(cbc_raw)
    if len(REQUIRED_FIELDS.intersection(cbc.keys())) < 3 and generic_full_text:
        cbc = _extract_cbc_from_text_lines(generic_full_text)

    # Validar que se encontraron suficientes campos minimos
    found_required = REQUIRED_FIELDS.intersection(cbc.keys())
    if len(found_required) < 3:
        raise ExtractionError(
            "No se encontraron suficientes valores del hemograma en el PDF. "
            f"Se esperaban al menos {REQUIRED_FIELDS}. "
            f"Se encontraron: {set(cbc.keys()) or 'ninguno'}."
        )

    # Resolver age_years desde el encabezado
    age_years = _parse_age_years(header.get("age"))

    # Resolver ubicacion desde el texto del PDF
    full_text_for_location = ""
    if not idexx_found:
        try:
            pdf3 = pdfplumber.open(io.BytesIO(pdf_bytes))
            full_text_for_location = "\n".join(
                (p.extract_text() or "") for p in pdf3.pages
            )
            pdf3.close()
        except Exception:
            pass
    else:
        full_text_for_location = header.get("clinic") or ""

    location = _extract_location(full_text_for_location, header.get("clinic"))

    metadata: dict[str, Optional[str]] = {
        "patient_name": header.get("patient_name"),
        "pet_owner": header.get("pet_owner"),
        "clinic": header.get("clinic"),
        "species": header.get("species"),
        "breed": header.get("breed"),
        "gender": header.get("gender"),
        "date_receipt": header.get("date_receipt"),
        "date_result": header.get("date_result"),
        "age_str": header.get("age"),
        "age_years": str(age_years) if age_years is not None else None,
        "location": location,
    }

    # Attach age_years into cbc dict so feature builder can use it directly
    if age_years is not None:
        cbc["age_years"] = age_years

    return ExtractionResult(cbc=cbc, metadata=metadata, comments=comments)


# ---------------------------------------------------------------------------
# Extraccion CSV / Excel
# ---------------------------------------------------------------------------

# Alias comunes de encabezados de columna (en minusculas, sin espacios) mapeados a nombres canonicos CBC.
# Se cubren tres fuentes reales del proyecto:
#   1. DAP raw  : prefijo krt_cbc_* (p.ej. krt_cbc_total_wbcs, krt_cbc_abs_neutrophils)
#   2. DAP parsed (tras NB02): columnas limpias, sufijo _pct para diferenciales relativos
#   3. IDEXX CSV (tras NB01) : Hematocrit, Hemoglobin, pct_Neutrophils, Plateletcrit, etc.
_CSV_ALIASES: dict[str, str] = {
    # --- WBC ---
    "wbc": "WBC",
    "leucocitos": "WBC",
    "krt_cbc_total_wbcs": "WBC",
    # --- RBC ---
    "rbc": "RBC",
    "eritrocitos": "RBC",
    "krt_cbc_rbc": "RBC",
    # --- HGB ---
    "hgb": "HGB",
    "hb": "HGB",
    "hemoglobin": "HGB",
    "hemoglobina": "HGB",
    "krt_cbc_hgb": "HGB",
    # --- HCT ---
    "hct": "HCT",
    "ht": "HCT",
    "hematocrit": "HCT",
    "hematocrito": "HCT",
    "pcv": "HCT",
    "krt_cbc_hct": "HCT",
    "krt_cbc_pcv": "HCT",
    # --- MCV ---
    "mcv": "MCV",
    "krt_cbc_mcv": "MCV",
    # --- MCH ---
    "mch": "MCH",
    "krt_cbc_mch": "MCH",
    # --- MCHC ---
    "mchc": "MCHC",
    "krt_cbc_mchc": "MCHC",
    # --- RDW ---
    "rdw": "RDW",
    "krt_cbc_rdw": "RDW",
    # --- Plaquetas (absolutas) ---
    "plt": "Platelets",
    "platelets": "Platelets",
    "plaquetas": "Platelets",
    "thrombocytes": "Platelets",
    "krt_cbc_plt": "Platelets",
    # --- MPV ---
    "mpv": "MPV",
    "krt_cbc_mpv": "MPV",
    # --- PDW ---
    "pdw": "PDW",
    # --- Plateletcrit ---
    "plateletcrit": "PCT",
    "krt_cbc_pct": "PCT",
    # --- Reticulocitos ---
    "reticulocytes": "Reticulocytes",
    "krt_cbc_retic_abs": "Reticulocytes",
    "pct_reticulocytes": "Reticulocytes_pct",
    "reticulocytes_pct": "Reticulocytes_pct",
    "krt_cbc_retic_per": "Reticulocytes_pct",
    # --- Neutrofilos (absolutos) ---
    "neu": "Neutrophils",
    "neut": "Neutrophils",
    "neutrophils": "Neutrophils",
    "neutrofilos": "Neutrophils",
    "seg": "Neutrophils",
    "krt_cbc_abs_neutrophils": "Neutrophils",
    # --- Neutrofilos (relativos, %) ---
    "pct_neutrophils": "Neutrophils_pct",
    "neutrophils_pct": "Neutrophils_pct",
    "krt_cbc_rel_neutrophils": "Neutrophils_pct",
    # --- Linfocitos (absolutos) ---
    "lym": "Lymphocytes",
    "lymph": "Lymphocytes",
    "lymphocytes": "Lymphocytes",
    "linfocitos": "Lymphocytes",
    "krt_cbc_abs_lymphocytes": "Lymphocytes",
    # --- Linfocitos (relativos, %) ---
    "pct_lymphocytes": "Lymphocytes_pct",
    "lymphocytes_pct": "Lymphocytes_pct",
    "krt_cbc_rel_lymphocytes": "Lymphocytes_pct",
    # --- Monocitos (absolutos) ---
    "mono": "Monocytes",
    "monocytes": "Monocytes",
    "monocitos": "Monocytes",
    "krt_cbc_abs_monocytes": "Monocytes",
    # --- Monocitos (relativos, %) ---
    "pct_monocytes": "Monocytes_pct",
    "monocytes_pct": "Monocytes_pct",
    "krt_cbc_rel_monocytes": "Monocytes_pct",
    # --- Eosinofilos (absolutos) ---
    "eos": "Eosinophils",
    "eosinophils": "Eosinophils",
    "eosinofilos": "Eosinophils",
    "krt_cbc_abs_eosinophils": "Eosinophils",
    # --- Eosinofilos (relativos, %) ---
    "pct_eosinophils": "Eosinophils_pct",
    "eosinophils_pct": "Eosinophils_pct",
    "krt_cbc_rel_eosinophils": "Eosinophils_pct",
    # --- Basofilos (absolutos) ---
    "baso": "Basophils",
    "basophils": "Basophils",
    "basofilos": "Basophils",
    "krt_cbc_abs_basophils": "Basophils",
    # --- Basofilos (relativos, %) ---
    "pct_basophils": "Basophils_pct",
    "basophils_pct": "Basophils_pct",
    "krt_cbc_rel_basophils": "Basophils_pct",
}


def _parse_tabular_df(df) -> dict[str, float]:  # type: ignore[type-arg]
    """
    Convierte un DataFrame de pandas en un diccionario CBC canonico.

    Soporta dos layouts habituales:
    1. Ancho (una fila por muestra): los encabezados de columna son los parametros, la primera fila tiene los valores.
    2. Clave-valor: dos columnas donde la columna 0 = nombre del parametro, columna 1 = valor numerico.
    """

    cbc: dict[str, float] = {}

    # Normalizar nombres de columna a cadenas simples
    df.columns = [str(c).strip() for c in df.columns]

    # Detectar layout clave-valor: las entradas de la primera columna coinciden con nombres CBC conocidos
    if df.shape[1] >= 2:
        first_col_lower = [str(v).strip().lower() for v in df.iloc[:, 0].dropna()]
        known_hits = sum(1 for v in first_col_lower if v in _CSV_ALIASES)
        if known_hits >= 2:
            for _, row in df.iterrows():
                key = str(row.iloc[0]).strip().lower()
                canonical = _CSV_ALIASES.get(key)
                if canonical:
                    parsed_value = _coerce_lab_number(row.iloc[1])
                    if parsed_value is not None:
                        cbc[canonical] = parsed_value
            return cbc

    # Layout ancho: los encabezados de columna son los nombres de los parametros
    if df.empty:
        return cbc
    first_row = df.iloc[0]
    for col in df.columns:
        canonical = _CSV_ALIASES.get(col.lower().strip())
        if canonical is None:
            continue
        parsed_value = _coerce_lab_number(first_row[col])
        if parsed_value is not None:
            cbc[canonical] = parsed_value

    return cbc


def _extract_from_csv(contents: bytes) -> ExtractionResult:
    """
    Extrae valores CBC de un archivo CSV.

    Soporta layouts en filas (campo, valor) y en columnas (encabezados = parametros).
    """
    import io as _io  # noqa: PLC0415

    import pandas as pd  # noqa: PLC0415

    df = None
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        for sep in (None, ",", ";", "\t"):
            try:
                read_kwargs = {"encoding": enc}
                if sep is None:
                    read_kwargs.update(sep=None, engine="python")
                else:
                    read_kwargs["sep"] = sep
                df = pd.read_csv(_io.BytesIO(contents), **read_kwargs)
                if df.shape[1] > 1 or sep is not None:
                    break
            except UnicodeDecodeError:
                continue
            except Exception:
                continue
        if df is not None:
            break

    if df is None:
        raise ExtractionError(
            "No se puede decodificar el CSV. Verifica la codificacion del archivo."
        )

    cbc = _parse_tabular_df(df)
    if not cbc:
        cbc = _extract_cbc_from_text_lines(decode_text(contents))

    found_required = REQUIRED_FIELDS.intersection(cbc.keys())
    if len(found_required) < 3:
        raise ExtractionError(
            "No se encontraron suficientes valores del hemograma en el CSV. "
            f"Se esperaban columnas como: {', '.join(sorted(REQUIRED_FIELDS))}. "
            f"Encontradas: {', '.join(sorted(cbc.keys())) or 'ninguna'}."
        )

    return ExtractionResult(cbc=cbc, metadata={"species": "Canino"}, comments=None)


def _extract_from_excel(contents: bytes) -> ExtractionResult:
    """
    Extrae valores CBC de un archivo Excel (.xlsx / .xls).
    """
    import io as _io  # noqa: PLC0415

    import pandas as pd  # noqa: PLC0415

    try:
        sheets = pd.read_excel(
            _io.BytesIO(contents), sheet_name=None, engine="openpyxl"
        )
    except Exception as exc:
        raise ExtractionError(f"No se pudo leer el Excel: {exc}") from exc

    cbc: dict[str, float] = {}
    text_parts: list[str] = []
    for sheet_name, df in sheets.items():
        cbc.update(_parse_tabular_df(df))
        text_parts.append(f"Sheet: {sheet_name}")
        text_parts.append(df.to_csv(index=False))
    if not cbc:
        cbc = _extract_cbc_from_text_lines("\n".join(text_parts))

    found_required = REQUIRED_FIELDS.intersection(cbc.keys())
    if len(found_required) < 3:
        raise ExtractionError(
            "No se encontraron suficientes valores del hemograma en el Excel. "
            f"Se esperaban columnas como: {', '.join(sorted(REQUIRED_FIELDS))}. "
            f"Encontradas: {', '.join(sorted(cbc.keys())) or 'ninguna'}."
        )

    return ExtractionResult(cbc=cbc, metadata={"species": "Canino"}, comments=None)


def _extract_from_txt(contents: bytes) -> ExtractionResult:
    text = decode_text(contents)
    cbc = _extract_cbc_from_text_lines(text)
    found_required = REQUIRED_FIELDS.intersection(cbc.keys())
    if len(found_required) < 3:
        raise ExtractionError(
            "No se encontraron suficientes valores del hemograma en el texto. "
            f"Se esperaban campos como: {', '.join(sorted(REQUIRED_FIELDS))}. "
            f"Encontrados: {', '.join(sorted(cbc.keys())) or 'ninguno'}."
        )
    return ExtractionResult(cbc=cbc, metadata={"species": "Canino"}, comments=None)


def _clean_ocr_text(text: str) -> str:
    """
    Limpia el texto bruto producido por Tesseract para mejorar el parseo CBC.

    Correcciones aplicadas:
    - Elimina lineas de un solo caracter (basura tipica de OCR).
    - Normaliza multiples espacios en blanco.
    - Corrige confusion O/0 en contextos numericos (p.ej. "8.O" -> "8.0").
    """
    import re as _re

    lines = text.split("\n")
    cleaned: list[str] = []
    for line in lines:
        stripped = line.strip()
        # Descartar lineas basura de 1 caracter
        if len(stripped) <= 1:
            continue
        # Normalizar whitespace interno
        stripped = _re.sub(r" {2,}", " ", stripped)
        # Corregir "digito.O" -> "digito.0" (O/0 confundidos por OCR)
        stripped = _re.sub(r"(\d)\.O(\b)", r"\g<1>.0\2", stripped)
        cleaned.append(stripped)
    return "\n".join(cleaned)


def _ocr_image_text_variants(img, pytesseract) -> list[str]:  # type: ignore[no-untyped-def]
    """Run OCR on the full image and likely table crops for photographed screens."""
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps  # noqa: PLC0415

    variants: list[str] = []

    def _prepare(candidate):  # type: ignore[no-untyped-def]
        candidate = ImageOps.autocontrast(candidate.convert("L"))
        candidate = ImageEnhance.Contrast(candidate).enhance(2.0)
        if candidate.width < 1000:
            factor = 1000 / max(candidate.width, 1)
            candidate = candidate.resize(
                (1000, int(candidate.height * factor)),
                Image.LANCZOS,
            )
        return candidate.filter(ImageFilter.SHARPEN)

    candidates = [img]
    width, height = img.size
    if width > height:
        crop_boxes = (
            (
                int(width * 0.25),
                int(height * 0.12),
                int(width * 0.70),
                int(height * 0.72),
            ),
            (
                int(width * 0.28),
                int(height * 0.15),
                int(width * 0.65),
                int(height * 0.66),
            ),
        )
        candidates.extend(img.crop(box) for box in crop_boxes)

    for idx, candidate in enumerate(candidates):
        prepared = _prepare(candidate)
        if idx > 0 and prepared.width < 1800:
            factor = 1800 / max(prepared.width, 1)
            prepared = prepared.resize(
                (1800, int(prepared.height * factor)),
                Image.LANCZOS,
            )
        for lang, psm in (("spa+eng", 6), ("eng", 6), ("eng", 4)):
            try:
                text = pytesseract.image_to_string(
                    prepared,
                    lang=lang,
                    config=f"--psm {psm} --oem 3",
                    timeout=10,
                ).strip()
            except Exception:
                continue
            if text:
                variants.append(text)
    return variants


def _extract_from_image(contents: bytes) -> ExtractionResult:
    """
    Extrae valores CBC de una imagen de hemograma usando OCR.

    Pipeline:
    1. Abrir con Pillow -> escala de grises -> autocontraste
    2. Escalar a minimo 1000px si la imagen es pequena
    3. MedianFilter para reducir ruido
    4. OCR con pytesseract (spa+eng, psm 6, oem 3)
    5. Limpiar texto y reutilizar parsers existentes
    6. Lanza ExtractionError si no se encuentran suficientes campos CBC
    """
    try:
        from PIL import Image, ImageFilter, ImageOps  # noqa: PLC0415
    except ImportError:
        raise ExtractionError(
            "Pillow no esta instalado. Contacta al administrador del servidor."
        )

    try:
        import pytesseract  # noqa: PLC0415
        from pytesseract import TesseractNotFoundError  # noqa: PLC0415
    except ImportError:
        raise ExtractionError(
            "pytesseract no esta instalado. Contacta al administrador del servidor."
        )

    # 1. Abrir y preprocesar la imagen
    try:
        img = Image.open(io.BytesIO(contents)).convert("L")
    except Exception as exc:
        raise ExtractionError(f"Formato de imagen no reconocido: {exc}") from exc

    img = ImageOps.autocontrast(img)

    # Escalar a minimo 1000px de ancho si la imagen es pequena
    if img.width < 1000:
        factor = 1000 / img.width
        new_size = (1000, int(img.height * factor))
        img = img.resize(new_size, Image.LANCZOS)

    img = img.filter(ImageFilter.MedianFilter(size=3))

    # 2. OCR
    try:
        raw_variants = _ocr_image_text_variants(img, pytesseract)
        raw_text = "\n".join(raw_variants)
    except TesseractNotFoundError:
        raise ExtractionError(
            "OCR no disponible en el servidor. Tesseract no esta instalado."
        )
    except Exception as exc:
        raise ExtractionError(f"Error durante OCR: {exc}") from exc

    # 3. Limpiar texto
    cleaned = _clean_ocr_text(raw_text)

    # 4. Intentar parsers existentes: IDEXX primero, luego generico
    cbc_raw = _extract_cbc_idexx(cleaned)
    header = _extract_header(cleaned) if cbc_raw else {}

    found_required = REQUIRED_FIELDS.intersection(_rename_to_canonical(cbc_raw).keys())
    if len(found_required) < 3:
        cbc_raw = _extract_cbc_generic(cleaned)
        header = {}

    cbc = _rename_to_canonical(cbc_raw)
    if len(REQUIRED_FIELDS.intersection(cbc.keys())) < 3:
        cbc = _extract_cbc_from_text_lines(cleaned)
    found_required = REQUIRED_FIELDS.intersection(cbc.keys())
    if len(found_required) < 3:
        raise ExtractionError(
            "No se encontraron suficientes valores del hemograma en la imagen. "
            "Por favor usa una imagen de mayor resolucion o mejor calidad. "
            f"Campos requeridos: {REQUIRED_FIELDS}. "
            f"Encontrados: {set(cbc.keys()) or 'ninguno'}."
        )

    # 5. Construir metadata y retornar
    age_years = _parse_age_years(header.get("age"))
    location = _extract_location(cleaned, header.get("clinic"))

    metadata: dict[str, Optional[str]] = {
        "patient_name": header.get("patient_name"),
        "pet_owner": header.get("pet_owner"),
        "clinic": header.get("clinic"),
        "species": header.get("species") or "Canino",
        "breed": header.get("breed"),
        "gender": header.get("gender"),
        "date_result": header.get("date_result"),
        "age_str": header.get("age"),
        "age_years": str(age_years) if age_years is not None else None,
        "location": location,
    }

    if age_years is not None:
        cbc["age_years"] = age_years

    return ExtractionResult(cbc=cbc, metadata=metadata, comments=None)


def extract_from_file(
    contents: bytes,
    content_type: str,
    filename: str | None = None,
) -> ExtractionResult:
    """
    Dispatcher publico: extrae valores CBC del archivo segun su tipo MIME.

    Soporta PDF, CSV, Excel (.xlsx/.xls) e imagenes (JPG, PNG, TIFF, WEBP).
    Si el tipo no se reconoce, intenta tratarlo como PDF (compatibilidad).
    """
    ct = (content_type or "").lower()
    ext = (
        (filename or "").lower().rsplit(".", 1)[-1]
        if filename and "." in filename
        else ""
    )
    if "csv" in ct or ct == "text/tab-separated-values" or ext in {"csv", "tsv"}:
        return _extract_from_csv(contents)
    if ct == "text/plain" or ext == "txt":
        return _extract_from_txt(contents)
    if (
        "spreadsheet" in ct
        or "excel" in ct
        or "xlsx" in ct
        or "xls" in ct
        or ext in {"xlsx", "xls"}
    ):
        return _extract_from_excel(contents)
    if ext == "docx" or "wordprocessingml.document" in ct:
        text = _extract_docx_text(contents)
        return _extract_from_txt(text.encode("utf-8"))
    if any(
        img_type in ct
        for img_type in ("image/jpeg", "image/png", "image/tiff", "image/webp")
    ):
        return _extract_from_image(contents)
    return extract_from_pdf(contents)


def _extract_docx_text(contents: bytes) -> str:
    try:
        from docx import Document  # noqa: PLC0415
    except Exception as exc:
        raise ExtractionError(
            "python-docx no esta instalado. Contacta al administrador del servidor."
        ) from exc

    try:
        document = Document(io.BytesIO(contents))
    except Exception as exc:
        raise ExtractionError(f"No se pudo leer el DOCX: {exc}") from exc

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)
