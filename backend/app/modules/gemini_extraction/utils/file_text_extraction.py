"""Convert supported upload formats into plain text for LLM extraction."""

from __future__ import annotations

import io
from pathlib import Path

from app.modules.gemini_extraction.schemas import ExtractionAttemptError


def decode_text(contents: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return contents.decode(encoding)
        except UnicodeDecodeError:
            continue
    return contents.decode("utf-8", errors="replace")


def extension(filename: str | None) -> str:
    return Path(filename or "").suffix.lower()


def is_pdf(content_type: str, filename: str | None) -> bool:
    return (
        extension(filename) == ".pdf"
        or (content_type or "").lower() == "application/pdf"
    )


def is_image(content_type: str, filename: str | None) -> bool:
    ext = extension(filename)
    ct = (content_type or "").lower()
    return ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"} or ct.startswith(
        "image/"
    )


def is_excel(content_type: str, filename: str | None) -> bool:
    ext = extension(filename)
    ct = (content_type or "").lower()
    return ext in {".xls", ".xlsx"} or "spreadsheet" in ct or "excel" in ct


def is_csv(content_type: str, filename: str | None) -> bool:
    ext = extension(filename)
    ct = (content_type or "").lower()
    return ext in {".csv", ".tsv"} or ct in {"text/csv", "text/tab-separated-values"}


def is_txt(content_type: str, filename: str | None) -> bool:
    ext = extension(filename)
    ct = (content_type or "").lower()
    return ext == ".txt" or ct == "text/plain"


def is_docx(content_type: str, filename: str | None) -> bool:
    ext = extension(filename)
    ct = (content_type or "").lower()
    return ext == ".docx" or "wordprocessingml.document" in ct


def _pdf_to_text(contents: bytes) -> str:
    try:
        import pdfplumber  # noqa: PLC0415
    except Exception as exc:
        raise ExtractionAttemptError(
            "PDF_DEPENDENCY_MISSING",
            "pdfplumber no esta instalado.",
        ) from exc

    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"Page {page_number}\n{text.strip()}")
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []
                for table in tables:
                    for row in table:
                        cells = [str(cell or "").strip() for cell in row]
                        if any(cells):
                            parts.append(" | ".join(cells))
                if len(text.split()) < 25:
                    ocr_text = _ocr_pdf_page(page)
                    if ocr_text:
                        parts.append(f"OCR Page {page_number}\n{ocr_text}")
    except Exception as exc:
        raise ExtractionAttemptError(
            "PDF_TEXT_ERROR", f"No se pudo leer el PDF: {exc}"
        ) from exc

    return "\n".join(parts).strip()


def _ocr_pdf_page(page: object) -> str:
    try:
        import pytesseract  # noqa: PLC0415
    except Exception:
        return ""
    try:
        image = page.to_image(resolution=200).original  # type: ignore[attr-defined]
        return pytesseract.image_to_string(
            image,
            lang="spa+eng",
            config="--psm 6 --oem 3",
            timeout=10,
        ).strip()
    except Exception:
        return ""


def _image_to_text(contents: bytes) -> str:
    try:
        from PIL import Image, ImageFilter, ImageOps  # noqa: PLC0415
        import pytesseract  # noqa: PLC0415
    except Exception as exc:
        raise ExtractionAttemptError(
            "OCR_DEPENDENCY_MISSING",
            "Pillow o pytesseract no estan instalados.",
        ) from exc

    try:
        image = Image.open(io.BytesIO(contents)).convert("L")
    except Exception as exc:
        raise ExtractionAttemptError(
            "IMAGE_READ_ERROR", f"No se pudo leer la imagen: {exc}"
        ) from exc

    image = ImageOps.autocontrast(image)
    if image.width < 1000:
        factor = 1000 / max(image.width, 1)
        image = image.resize((1000, int(image.height * factor)), Image.LANCZOS)
    image = image.filter(ImageFilter.MedianFilter(size=3))
    try:
        return pytesseract.image_to_string(
            image,
            lang="spa+eng",
            config="--psm 6 --oem 3",
            timeout=10,
        ).strip()
    except Exception as exc:
        raise ExtractionAttemptError(
            "OCR_ERROR", f"No se pudo ejecutar OCR: {exc}"
        ) from exc


def _excel_to_text(contents: bytes) -> str:
    try:
        import pandas as pd  # noqa: PLC0415
    except Exception as exc:
        raise ExtractionAttemptError(
            "PANDAS_MISSING", "pandas no esta instalado."
        ) from exc

    try:
        sheets = pd.read_excel(io.BytesIO(contents), sheet_name=None)
    except Exception as exc:
        raise ExtractionAttemptError(
            "EXCEL_READ_ERROR", f"No se pudo leer el Excel: {exc}"
        ) from exc

    parts: list[str] = []
    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        parts.append(df.to_csv(index=False))
    return "\n".join(parts).strip()


def _csv_to_text(contents: bytes) -> str:
    text = decode_text(contents)
    return text.strip()


def _docx_to_text(contents: bytes) -> str:
    try:
        from docx import Document  # noqa: PLC0415
    except Exception as exc:
        raise ExtractionAttemptError(
            "DOCX_DEPENDENCY_MISSING",
            "python-docx no esta instalado.",
        ) from exc

    try:
        document = Document(io.BytesIO(contents))
    except Exception as exc:
        raise ExtractionAttemptError(
            "DOCX_READ_ERROR", f"No se pudo leer el DOCX: {exc}"
        ) from exc

    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_file(
    *,
    contents: bytes,
    content_type: str,
    filename: str | None,
) -> str:
    if is_pdf(content_type, filename):
        return _pdf_to_text(contents)
    if is_image(content_type, filename):
        return _image_to_text(contents)
    if is_excel(content_type, filename):
        return _excel_to_text(contents)
    if is_csv(content_type, filename) or is_txt(content_type, filename):
        return _csv_to_text(contents)
    if is_docx(content_type, filename):
        return _docx_to_text(contents)
    return decode_text(contents).strip()
